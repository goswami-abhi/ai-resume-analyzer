import os
import pdfplumber
import PyPDF2
from docx import Document

def extract_text_from_pdf(file_path):
    """
    Extracts all text from a PDF file using pdfplumber, with a fallback to PyPDF2.
    """
    text = ""
    # Try pdfplumber first
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed: {e}. Trying PyPDF2...")
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e2:
            print(f"PyPDF2 failed: {e2}")
            raise Exception(f"Could not parse PDF file: {str(e2)}")
            
    return text.strip()

def extract_text_from_docx(file_path):
    """
    Extracts all text from a DOCX file using python-docx.
    """
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text).strip()
    except Exception as e:
        print(f"python-docx failed: {e}")
        raise Exception(f"Could not parse DOCX file: {str(e)}")

def extract_resume_text(file_path):
    """
    Determines file type and extracts text accordingly.
    """
    _, ext = os.path.splitext(file_path.lower())
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX files are allowed.")
